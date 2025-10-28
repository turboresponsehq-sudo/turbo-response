from flask import Blueprint, request, jsonify
import json
import os
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import mimetypes

automation_bp = Blueprint('automation', __name__)

# ============================================================================
# AUTOMATION CONFIGURATION
# ============================================================================

AUTOMATION_CONFIG = {
    'notification_email': 'TurboResponseHQ@gmail.com',
    'data_storage_path': '/data/turbo_response_data',
    'document_storage_path': '/data/turbo_response_documents',
    'enable_notifications': True,
    'enable_data_logging': True,
    'enable_email_sending': True,  # Set to True when SMTP is configured
    'smtp_server': os.getenv('SMTP_SERVER', 'smtp.gmail.com'),
    'smtp_port': int(os.getenv('SMTP_PORT', 587)),
    'smtp_username': os.getenv('SMTP_USERNAME', ''),
    'smtp_password': os.getenv('SMTP_PASSWORD', ''),
    'sender_email': os.getenv('SENDER_EMAIL', 'TurboResponseHQ@gmail.com'),
}

# Ensure directories exist
os.makedirs(AUTOMATION_CONFIG['data_storage_path'], exist_ok=True)
os.makedirs(AUTOMATION_CONFIG['document_storage_path'], exist_ok=True)

# ============================================================================
# FORM SUBMISSION HANDLER
# ============================================================================

@automation_bp.route('/api/form-submission', methods=['POST'])
def handle_form_submission():
    """
    Handle intake form submissions with full automation:
    1. Store in database
    2. Send admin notification
    3. Send client confirmation
    4. Handle document uploads
    5. Log for external integrations
    """
    try:
        data = request.get_json()
        
        # Generate case ID
        case_id = f"TURBO-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
        
        # Create submission record
        submission_data = {
            'case_id': case_id,
            'timestamp': datetime.now().isoformat(),
            'client_data': {
                'email': data.get('email'),
                'fullName': data.get('fullName'),
                'phone': data.get('phone'),
                'address': data.get('address'),
                'category': data.get('category'),
                'caseDescription': data.get('caseDescription'),
                'amount': data.get('amount'),
                'deadline': data.get('deadline'),
                'documents': data.get('documents', [])
            },
            'contract_acceptance': {
                'agreed': data.get('contractAgreement', False),
                'signature': data.get('clientSignature', ''),
                'signature_date': data.get('signatureDate', ''),
                'ip_address': request.remote_addr,
                'timestamp': datetime.now().isoformat()
            },
            'status': 'submitted',
            'automation_triggers': []
        }
        
        # Trigger 1: Store submission in database
        if AUTOMATION_CONFIG['enable_data_logging']:
            store_submission_data(submission_data)
            submission_data['automation_triggers'].append('data_stored')
        
        # Trigger 2: Handle document uploads
        if data.get('documents'):
            handle_document_uploads(case_id, data.get('documents'))
            submission_data['automation_triggers'].append('documents_stored')
        
        # Trigger 3: Send admin notification email
        if AUTOMATION_CONFIG['enable_notifications']:
            send_admin_notification(submission_data)
            submission_data['automation_triggers'].append('admin_notified')
        
        # Trigger 4: Send client confirmation email
        send_client_confirmation(submission_data)
        submission_data['automation_triggers'].append('client_confirmed')
        
        # Trigger 5: Log for external integrations (Zapier/Notion ready)
        log_for_external_integration(submission_data)
        submission_data['automation_triggers'].append('external_logged')
        
        return jsonify({
            'success': True,
            'case_id': case_id,
            'message': 'Case submitted successfully. You will receive confirmation within 24-48 hours.',
            'automation_status': submission_data['automation_triggers']
        }), 200
        
    except Exception as e:
        print(f"Form submission error: {e}")
        return jsonify({
            'success': False,
            'error': str(e),
            'message': 'Submission failed. Please try again.'
        }), 500

# ============================================================================
# CHATBOT LEAD HANDLER
# ============================================================================

@automation_bp.route('/api/chatbot-lead', methods=['POST'])
def handle_chatbot_lead():
    """Handle chatbot lead capture with automation"""
    try:
        lead_data = request.get_json()
        
        # Add metadata
        lead_record = {
            'lead_id': f"LEAD-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
            'timestamp': datetime.now().isoformat(),
            'source': 'turbo_ai_chatbot',
            'lead_data': lead_data,
            'status': 'new_lead'
        }
        
        # Store lead data
        store_lead_data(lead_record)
        
        # Send lead notification
        send_lead_notification(lead_record)
        
        # Log for CRM integration
        log_lead_for_crm(lead_record)
        
        return jsonify({
            'success': True,
            'lead_id': lead_record['lead_id'],
            'message': 'Lead captured successfully'
        }), 200
        
    except Exception as e:
        print(f"Chatbot lead error: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ============================================================================
# WEBHOOK ENDPOINT FOR ZAPIER
# ============================================================================

@automation_bp.route('/api/webhook/zapier', methods=['POST'])
def zapier_webhook():
    """
    Webhook endpoint for Zapier integrations
    
    Accepts POST requests with form field data and triggers automations
    """
    try:
        data = request.get_json()
        
        # Process Zapier webhook data
        webhook_record = {
            'webhook_id': f"ZAP-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
            'timestamp': datetime.now().isoformat(),
            'source': 'zapier',
            'data': data
        }
        
        # Store webhook data
        store_webhook_data(webhook_record)
        
        return jsonify({
            'success': True,
            'webhook_id': webhook_record['webhook_id']
        }), 200
        
    except Exception as e:
        print(f"Zapier webhook error: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ============================================================================
# DATA STORAGE FUNCTIONS
# ============================================================================

def store_submission_data(submission_data):
    """Store form submission data to local filesystem and database"""
    try:
        # Store individual case file
        filename = f"{submission_data['case_id']}_submission.json"
        filepath = os.path.join(AUTOMATION_CONFIG['data_storage_path'], filename)
        
        with open(filepath, 'w') as f:
            json.dump(submission_data, f, indent=2)
        
        # Append to master log
        log_filepath = os.path.join(AUTOMATION_CONFIG['data_storage_path'], 'submissions_log.jsonl')
        with open(log_filepath, 'a') as f:
            f.write(json.dumps(submission_data) + '\n')
        
        print(f"✓ Submission stored: {submission_data['case_id']}")
        
    except Exception as e:
        print(f"✗ Error storing submission data: {e}")

def store_lead_data(lead_record):
    """Store chatbot lead data"""
    try:
        filename = f"{lead_record['lead_id']}_lead.json"
        filepath = os.path.join(AUTOMATION_CONFIG['data_storage_path'], filename)
        
        with open(filepath, 'w') as f:
            json.dump(lead_record, f, indent=2)
        
        # Append to leads log
        log_filepath = os.path.join(AUTOMATION_CONFIG['data_storage_path'], 'leads_log.jsonl')
        with open(log_filepath, 'a') as f:
            f.write(json.dumps(lead_record) + '\n')
        
        print(f"✓ Lead stored: {lead_record['lead_id']}")
        
    except Exception as e:
        print(f"✗ Error storing lead data: {e}")

def store_webhook_data(webhook_record):
    """Store webhook data"""
    try:
        filename = f"{webhook_record['webhook_id']}_webhook.json"
        filepath = os.path.join(AUTOMATION_CONFIG['data_storage_path'], filename)
        
        with open(filepath, 'w') as f:
            json.dump(webhook_record, f, indent=2)
        
        print(f"✓ Webhook stored: {webhook_record['webhook_id']}")
        
    except Exception as e:
        print(f"✗ Error storing webhook data: {e}")

# ============================================================================
# DOCUMENT UPLOAD HANDLER
# ============================================================================

def handle_document_uploads(case_id, documents):
    """
    Handle document uploads and storage
    
    Documents can be:
    - Base64 encoded files from form
    - File paths for server-side processing
    """
    try:
        case_folder = os.path.join(AUTOMATION_CONFIG['document_storage_path'], case_id)
        os.makedirs(case_folder, exist_ok=True)
        
        for idx, doc in enumerate(documents):
            if isinstance(doc, dict):
                # Handle base64 encoded files
                if 'data' in doc and 'filename' in doc:
                    import base64
                    file_data = base64.b64decode(doc['data'])
                    file_path = os.path.join(case_folder, doc['filename'])
                    
                    with open(file_path, 'wb') as f:
                        f.write(file_data)
                    
                    print(f"✓ Document stored: {case_id}/{doc['filename']}")
        
        return True
        
    except Exception as e:
        print(f"✗ Error handling document uploads: {e}")
        return False

# ============================================================================
# EMAIL SENDING FUNCTIONS
# ============================================================================

def send_admin_notification(submission_data):
    """
    Send notification email to admin (TurboResponseHQ@gmail.com)
    
    Contains full case details for review
    """
    try:
        client_data = submission_data['client_data']
        
        subject = f"🔔 New Case: {submission_data['case_id']} - {client_data.get('fullName', 'Unknown')}"
        
        # Create HTML email body
        html_body = f"""
        <html>
            <body style="font-family: Arial, sans-serif; color: #333;">
                <h2 style="color: #06b6d4;">New Case Submission</h2>
                
                <div style="background: #f5f5f5; padding: 15px; border-radius: 8px; margin: 20px 0;">
                    <p><strong>Case ID:</strong> {submission_data['case_id']}</p>
                    <p><strong>Submitted:</strong> {datetime.fromisoformat(submission_data['timestamp']).strftime('%B %d, %Y at %I:%M %p')}</p>
                </div>
                
                <h3>Client Information</h3>
                <ul>
                    <li><strong>Name:</strong> {client_data.get('fullName', 'Not provided')}</li>
                    <li><strong>Email:</strong> {client_data.get('email', 'Not provided')}</li>
                    <li><strong>Phone:</strong> {client_data.get('phone', 'Not provided')}</li>
                    <li><strong>Address:</strong> {client_data.get('address', 'Not provided')}</li>
                </ul>
                
                <h3>Case Details</h3>
                <ul>
                    <li><strong>Category:</strong> {client_data.get('category', 'Not specified')}</li>
                    <li><strong>Amount:</strong> ${client_data.get('amount', 'Not specified')}</li>
                    <li><strong>Deadline:</strong> {client_data.get('deadline', 'Not specified')}</li>
                    <li><strong>Documents:</strong> {len(client_data.get('documents', []))} files</li>
                </ul>
                
                <h3>Issue Description</h3>
                <p style="background: #f9f9f9; padding: 10px; border-left: 4px solid #06b6d4;">
                    {client_data.get('caseDescription', 'No description provided')}
                </p>
                
                <hr style="margin: 20px 0; border: none; border-top: 1px solid #ddd;">
                
                <p>
                    <a href="https://turboresponsehq.ai/admin" 
                       style="background: #06b6d4; color: white; padding: 10px 20px; text-decoration: none; border-radius: 5px;">
                        Review in Admin Dashboard
                    </a>
                </p>
                
                <p style="color: #999; font-size: 12px; margin-top: 30px;">
                    This is an automated notification from Turbo Response
                </p>
            </body>
        </html>
        """
        
        # Log notification
        log_email_notification({
            'type': 'admin_notification',
            'to': AUTOMATION_CONFIG['notification_email'],
            'subject': subject,
            'case_id': submission_data['case_id'],
            'status': 'logged'
        })
        
        # Send actual email if SMTP is configured
        if AUTOMATION_CONFIG['enable_email_sending'] and AUTOMATION_CONFIG['smtp_username']:
            send_email(
                to_email=AUTOMATION_CONFIG['notification_email'],
                subject=subject,
                html_body=html_body,
                case_id=submission_data['case_id']
            )
        
    except Exception as e:
        print(f"✗ Error sending admin notification: {e}")

def send_client_confirmation(submission_data):
    """
    Send auto-reply confirmation email to client
    
    Provides case ID and next steps
    """
    try:
        client_data = submission_data['client_data']
        client_email = client_data.get('email')
        
        if not client_email:
            print("⚠ No client email provided")
            return
        
        subject = f"✅ Case Received: {submission_data['case_id']} - Turbo Response"
        
        html_body = f"""
        <html>
            <body style="font-family: Arial, sans-serif; color: #333;">
                <h2 style="color: #06b6d4;">✅ Your Case Has Been Received</h2>
                
                <p>Dear {client_data.get('fullName', 'Valued Client')},</p>
                
                <p>Thank you for submitting your case to Turbo Response. Our AI system is analyzing your information right now.</p>
                
                <div style="background: #f0f9ff; padding: 15px; border-radius: 8px; margin: 20px 0; border-left: 4px solid #06b6d4;">
                    <p><strong>Your Case ID:</strong> <span style="font-size: 18px; color: #06b6d4;">{submission_data['case_id']}</span></p>
                    <p><strong>Category:</strong> {client_data.get('category', 'Consumer Defense')}</p>
                    <p><strong>Submitted:</strong> {datetime.fromisoformat(submission_data['timestamp']).strftime('%B %d, %Y at %I:%M %p')}</p>
                </div>
                
                <h3>What Happens Next</h3>
                <ol>
                    <li><strong>AI Analysis</strong> - Our AI system analyzes your documents and case details</li>
                    <li><strong>Expert Review</strong> - Our consumer rights specialists review your situation</li>
                    <li><strong>Game Plan Creation</strong> - We create your personalized defense strategy</li>
                    <li><strong>Delivery</strong> - You receive your response within 24-48 hours</li>
                </ol>
                
                <h3>Keep Your Case ID Safe</h3>
                <p>You'll need your case ID to check on your case status. Save this email for your records.</p>
                
                <h3>Questions?</h3>
                <p>Reply to this email or visit our website at <a href="https://turboresponsehq.ai">turboresponsehq.ai</a></p>
                
                <hr style="margin: 20px 0; border: none; border-top: 1px solid #ddd;">
                
                <p style="color: #666;">
                    Best regards,<br>
                    <strong>The Turbo Response Team</strong><br>
                    AI-Powered Consumer Defense
                </p>
                
                <p style="color: #999; font-size: 12px;">
                    This is an automated confirmation. Please do not reply to this email.
                </p>
            </body>
        </html>
        """
        
        # Log client confirmation
        log_email_notification({
            'type': 'client_confirmation',
            'to': client_email,
            'subject': subject,
            'case_id': submission_data['case_id'],
            'status': 'logged'
        })
        
        # Send actual email if SMTP is configured
        if AUTOMATION_CONFIG['enable_email_sending'] and AUTOMATION_CONFIG['smtp_username']:
            send_email(
                to_email=client_email,
                subject=subject,
                html_body=html_body,
                case_id=submission_data['case_id']
            )
        
        print(f"✓ Client confirmation logged: {client_email}")
        
    except Exception as e:
        print(f"✗ Error sending client confirmation: {e}")

def send_lead_notification(lead_record):
    """Send lead notification email to admin"""
    try:
        lead_data = lead_record['lead_data']
        
        subject = f"💬 New Chatbot Lead: {lead_record['lead_id']}"
        
        html_body = f"""
        <html>
            <body style="font-family: Arial, sans-serif; color: #333;">
                <h2 style="color: #06b6d4;">New Lead from Turbo AI Chatbot</h2>
                
                <div style="background: #f5f5f5; padding: 15px; border-radius: 8px; margin: 20px 0;">
                    <p><strong>Lead ID:</strong> {lead_record['lead_id']}</p>
                    <p><strong>Captured:</strong> {datetime.fromisoformat(lead_record['timestamp']).strftime('%B %d, %Y at %I:%M %p')}</p>
                </div>
                
                <h3>Lead Information</h3>
                <ul>
                    <li><strong>Name:</strong> {lead_data.get('name', 'Not provided')}</li>
                    <li><strong>Email:</strong> {lead_data.get('email', 'Not provided')}</li>
                    <li><strong>Question:</strong> {lead_data.get('question', 'Not provided')}</li>
                </ul>
                
                <p style="color: #d97706; font-weight: bold;">⏰ Follow up within 24 hours for best conversion rates</p>
                
                <p style="color: #999; font-size: 12px; margin-top: 30px;">
                    This is an automated notification from Turbo Response
                </p>
            </body>
        </html>
        """
        
        # Log lead notification
        log_email_notification({
            'type': 'lead_notification',
            'to': AUTOMATION_CONFIG['notification_email'],
            'subject': subject,
            'lead_id': lead_record['lead_id'],
            'status': 'logged'
        })
        
        # Send actual email if SMTP is configured
        if AUTOMATION_CONFIG['enable_email_sending'] and AUTOMATION_CONFIG['smtp_username']:
            send_email(
                to_email=AUTOMATION_CONFIG['notification_email'],
                subject=subject,
                html_body=html_body
            )
        
    except Exception as e:
        print(f"✗ Error sending lead notification: {e}")

def send_email(to_email, subject, html_body, case_id=None):
    """
    Send email via SMTP
    
    Requires environment variables:
    - SMTP_SERVER (default: smtp.gmail.com)
    - SMTP_PORT (default: 587)
    - SMTP_USERNAME (your email)
    - SMTP_PASSWORD (app password)
    - SENDER_EMAIL (from address)
    """
    try:
        if not AUTOMATION_CONFIG['smtp_username'] or not AUTOMATION_CONFIG['smtp_password']:
            print("⚠ SMTP credentials not configured. Email not sent.")
            return False
        
        # Create message
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = AUTOMATION_CONFIG['sender_email']
        msg['To'] = to_email
        
        # Attach HTML body
        part = MIMEText(html_body, 'html')
        msg.attach(part)
        
        # Send email
        with smtplib.SMTP(AUTOMATION_CONFIG['smtp_server'], AUTOMATION_CONFIG['smtp_port']) as server:
            server.starttls()
            server.login(AUTOMATION_CONFIG['smtp_username'], AUTOMATION_CONFIG['smtp_password'])
            server.send_message(msg)
        
        print(f"✓ Email sent to {to_email}")
        return True
        
    except Exception as e:
        print(f"✗ Error sending email: {e}")
        return False

# ============================================================================
# LOGGING FUNCTIONS
# ============================================================================

def log_email_notification(notification):
    """Log email notification attempt"""
    try:
        log_filepath = os.path.join(AUTOMATION_CONFIG['data_storage_path'], 'email_notifications_log.jsonl')
        notification['timestamp'] = datetime.now().isoformat()
        
        with open(log_filepath, 'a') as f:
            f.write(json.dumps(notification) + '\n')
        
    except Exception as e:
        print(f"✗ Error logging email notification: {e}")

def log_for_external_integration(submission_data):
    """
    Log data in format ready for external integrations (Zapier, Notion, etc.)
    
    This creates a standardized format that can be easily consumed by:
    - Zapier webhooks
    - Notion API
    - Google Sheets
    - Custom integrations
    """
    try:
        external_data = {
            'case_id': submission_data['case_id'],
            'timestamp': submission_data['timestamp'],
            'client_name': submission_data['client_data'].get('fullName', ''),
            'client_email': submission_data['client_data'].get('email', ''),
            'client_phone': submission_data['client_data'].get('phone', ''),
            'client_address': submission_data['client_data'].get('address', ''),
            'category': submission_data['client_data'].get('category', ''),
            'issue_description': submission_data['client_data'].get('caseDescription', ''),
            'amount': submission_data['client_data'].get('amount', ''),
            'deadline': submission_data['client_data'].get('deadline', ''),
            'document_count': len(submission_data['client_data'].get('documents', [])),
            'status': 'new_submission',
            'integration_ready': True
        }
        
        # Store in integration-ready format
        log_filepath = os.path.join(AUTOMATION_CONFIG['data_storage_path'], 'external_integrations_log.jsonl')
        with open(log_filepath, 'a') as f:
            f.write(json.dumps(external_data) + '\n')
        
        print(f"✓ External integration data logged: {submission_data['case_id']}")
        
    except Exception as e:
        print(f"✗ Error logging for external integration: {e}")

def log_lead_for_crm(lead_record):
    """Log lead data for CRM integration"""
    try:
        crm_data = {
            'lead_id': lead_record['lead_id'],
            'timestamp': lead_record['timestamp'],
            'source': 'turbo_ai_chatbot',
            'name': lead_record['lead_data'].get('name', ''),
            'email': lead_record['lead_data'].get('email', ''),
            'question': lead_record['lead_data'].get('question', ''),
            'status': 'new_lead',
            'priority': 'medium',
            'follow_up_required': True
        }
        
        # Store in CRM-ready format
        log_filepath = os.path.join(AUTOMATION_CONFIG['data_storage_path'], 'crm_leads_log.jsonl')
        with open(log_filepath, 'a') as f:
            f.write(json.dumps(crm_data) + '\n')
        
    except Exception as e:
        print(f"✗ Error logging lead for CRM: {e}")

# ============================================================================
# STATUS & MONITORING
# ============================================================================

@automation_bp.route('/api/automation-status', methods=['GET'])
def get_automation_status():
    """Get automation system status and recent activity"""
    try:
        submissions_count = count_recent_files('submissions_log.jsonl')
        leads_count = count_recent_files('leads_log.jsonl')
        emails_logged = count_recent_files('email_notifications_log.jsonl')
        
        status = {
            'system_status': 'operational',
            'recent_submissions': submissions_count,
            'recent_leads': leads_count,
            'emails_logged': emails_logged,
            'automation_config': {
                'notifications_enabled': AUTOMATION_CONFIG['enable_notifications'],
                'data_logging_enabled': AUTOMATION_CONFIG['enable_data_logging'],
                'email_sending_enabled': AUTOMATION_CONFIG['enable_email_sending'],
                'smtp_configured': bool(AUTOMATION_CONFIG['smtp_username'])
            },
            'timestamp': datetime.now().isoformat()
        }
        
        return jsonify(status), 200
        
    except Exception as e:
        return jsonify({
            'system_status': 'error',
            'error': str(e)
        }), 500

def count_recent_files(filename):
    """Count recent entries in log file"""
    try:
        filepath = os.path.join(AUTOMATION_CONFIG['data_storage_path'], filename)
        if not os.path.exists(filepath):
            return 0
        
        with open(filepath, 'r') as f:
            lines = f.readlines()
            return len(lines)
        
    except Exception:
        return 0




# ============================================================================
# ADMIN API ENDPOINTS
# ============================================================================

@automation_bp.route('/api/admin/cases', methods=['GET'])
def get_all_cases():
    """Get all submitted cases for admin dashboard"""
    try:
        cases = []
        data_path = AUTOMATION_CONFIG['data_storage_path']
        
        # Read all case files
        if os.path.exists(data_path):
            for filename in os.listdir(data_path):
                if filename.endswith('_submission.json'):
                    filepath = os.path.join(data_path, filename)
                    try:
                        with open(filepath, 'r') as f:
                            case_data = json.load(f)
                            cases.append(case_data)
                    except Exception as e:
                        print(f"Error reading case file {filename}: {e}")
        
        # Sort by submission date (newest first)
        cases.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        
        return jsonify({
            'success': True,
            'cases': cases,
            'total': len(cases)
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@automation_bp.route('/api/admin/case/<case_id>', methods=['GET'])
def get_case_details(case_id):
    """Get details of a specific case"""
    try:
        data_path = AUTOMATION_CONFIG['data_storage_path']
        filepath = os.path.join(data_path, f"{case_id}_submission.json")
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'error': 'Case not found'
            }), 404
        
        with open(filepath, 'r') as f:
            case_data = json.load(f)
        
        return jsonify({
            'success': True,
            'case': case_data
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@automation_bp.route('/api/admin/document/<case_id>/<filename>', methods=['GET'])
def download_document(case_id, filename):
    """Download a document for a specific case"""
    try:
        doc_path = AUTOMATION_CONFIG['document_storage_path']
        case_doc_path = os.path.join(doc_path, case_id)
        filepath = os.path.join(case_doc_path, filename)
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'error': 'Document not found'
            }), 404
        
        # Determine MIME type
        mime_type, _ = mimetypes.guess_type(filename)
        if not mime_type:
            mime_type = 'application/octet-stream'
        
        from flask import send_file
        return send_file(filepath, mimetype=mime_type, as_attachment=True, download_name=filename)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@automation_bp.route('/api/admin/case/<case_id>/document/<int:doc_index>', methods=['GET'])
def download_document_by_index(case_id, doc_index):
    """Download a document by index for a specific case"""
    try:
        # Load case data to get document filename
        data_path = AUTOMATION_CONFIG['data_storage_path']
        case_file = os.path.join(data_path, f"{case_id}_submission.json")
        
        if not os.path.exists(case_file):
            return jsonify({
                'success': False,
                'error': 'Case not found'
            }), 404
        
        with open(case_file, 'r') as f:
            case_data = json.load(f)
        
        client_data = case_data.get('client_data', {})
        documents = client_data.get('documents', [])
        
        if doc_index >= len(documents):
            return jsonify({
                'success': False,
                'error': 'Document index out of range'
            }), 404
        
        doc = documents[doc_index]
        filename = doc.get('filename', f'document_{doc_index}')
        
        # Get document from storage
        doc_path = AUTOMATION_CONFIG['document_storage_path']
        case_doc_path = os.path.join(doc_path, case_id)
        filepath = os.path.join(case_doc_path, filename)
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'error': 'Document file not found on disk'
            }), 404
        
        # Determine MIME type
        mime_type, _ = mimetypes.guess_type(filename)
        if not mime_type:
            mime_type = 'application/octet-stream'
        
        from flask import send_file
        return send_file(filepath, mimetype=mime_type, as_attachment=True, download_name=filename)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@automation_bp.route('/api/admin/case/<case_id>/approve', methods=['POST'])
def approve_case_payment(case_id):
    """Mark a case payment as approved"""
    try:
        data_path = AUTOMATION_CONFIG['data_storage_path']
        filepath = os.path.join(data_path, f"{case_id}_submission.json")
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'error': 'Case not found'
            }), 404
        
        # Read case data
        with open(filepath, 'r') as f:
            case_data = json.load(f)
        
        # Update payment status
        case_data['payment_status'] = 'approved'
        case_data['payment_approved_at'] = datetime.now().isoformat()
        
        # Save updated data
        with open(filepath, 'w') as f:
            json.dump(case_data, f, indent=2)
        
        return jsonify({
            'success': True,
            'message': 'Payment approved',
            'case_id': case_id
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@automation_bp.route('/api/admin/case/<case_id>/status', methods=['POST'])
def update_case_status(case_id):
    """Update case status"""
    try:
        data = request.get_json()
        new_status = data.get('status')
        
        if not new_status:
            return jsonify({
                'success': False,
                'error': 'Status is required'
            }), 400
        
        data_path = AUTOMATION_CONFIG['data_storage_path']
        filepath = os.path.join(data_path, f"{case_id}_submission.json")
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'error': 'Case not found'
            }), 404
        
        # Read case data
        with open(filepath, 'r') as f:
            case_data = json.load(f)
        
        # Update status
        case_data['case_status'] = new_status
        case_data['status_updated_at'] = datetime.now().isoformat()
        
        # Save updated data
        with open(filepath, 'w') as f:
            json.dump(case_data, f, indent=2)
        
        return jsonify({
            'success': True,
            'message': 'Status updated',
            'case_id': case_id,
            'new_status': new_status
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500



# ============================================================================
# ADMIN AUTHENTICATION
# ============================================================================

import secrets
import hashlib
from functools import wraps

# Simple in-memory session storage (for production, use Redis or database)
admin_sessions = {}

# Admin credentials (in production, store hashed passwords in database)
ADMIN_CREDENTIALS = {
    'username': os.getenv('ADMIN_USERNAME', 'admin'),
    'password_hash': hashlib.sha256(os.getenv('ADMIN_PASSWORD', 'TurboAdmin2025!').encode()).hexdigest()
}

def verify_admin_session(f):
    """Decorator to verify admin session token"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        auth_header = request.headers.get('Authorization')
        
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({
                'success': False,
                'error': 'Unauthorized - No token provided'
            }), 401
        
        token = auth_header.split(' ')[1]
        
        if token not in admin_sessions:
            return jsonify({
                'success': False,
                'error': 'Unauthorized - Invalid or expired session'
            }), 401
        
        # Check if session is still valid (24 hour expiry)
        session_data = admin_sessions[token]
        session_age = (datetime.now() - session_data['created_at']).total_seconds()
        
        if session_age > 86400:  # 24 hours
            del admin_sessions[token]
            return jsonify({
                'success': False,
                'error': 'Unauthorized - Session expired'
            }), 401
        
        return f(*args, **kwargs)
    
    return decorated_function

@automation_bp.route('/api/admin/login', methods=['POST'])
def admin_login():
    """Admin login endpoint"""
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        
        if not username or not password:
            return jsonify({
                'success': False,
                'error': 'Username and password are required'
            }), 400
        
        # Verify credentials
        password_hash = hashlib.sha256(password.encode()).hexdigest()
        
        if username != ADMIN_CREDENTIALS['username'] or password_hash != ADMIN_CREDENTIALS['password_hash']:
            return jsonify({
                'success': False,
                'error': 'Invalid username or password'
            }), 401
        
        # Generate session token
        token = secrets.token_urlsafe(32)
        admin_sessions[token] = {
            'username': username,
            'created_at': datetime.now()
        }
        
        return jsonify({
            'success': True,
            'token': token,
            'message': 'Login successful'
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@automation_bp.route('/api/admin/verify', methods=['GET'])
@verify_admin_session
def verify_session():
    """Verify if admin session is still valid"""
    return jsonify({
        'success': True,
        'message': 'Session is valid'
    }), 200

@automation_bp.route('/api/admin/logout', methods=['POST'])
def admin_logout():
    """Admin logout endpoint"""
    try:
        auth_header = request.headers.get('Authorization')
        
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header.split(' ')[1]
            if token in admin_sessions:
                del admin_sessions[token]
        
        return jsonify({
            'success': True,
            'message': 'Logged out successfully'
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500




# ============================================================================
# PAYMENT LINK SYSTEM
# ============================================================================

# Store payment links in memory (in production, use database)
payment_links = {}

@automation_bp.route('/api/generate-payment-link', methods=['POST'])
def generate_payment_link():
    """Generate unique payment link for a case"""
    try:
        data = request.get_json()
        case_id = data.get('case_id')
        price = data.get('price')
        
        if not case_id or not price:
            return jsonify({'success': False, 'error': 'Missing case_id or price'}), 400
        
        # Load case data
        case_file = os.path.join(AUTOMATION_CONFIG['data_storage_path'], f"{case_id}_submission.json")
        if not os.path.exists(case_file):
            return jsonify({'success': False, 'error': 'Case not found'}), 404
        
        with open(case_file, 'r') as f:
            case_data = json.load(f)
        
        # Generate unique token
        import secrets
        payment_token = secrets.token_urlsafe(32)
        
        # Store payment link data
        payment_links[payment_token] = {
            'case_id': case_id,
            'price': price,
            'created_at': datetime.now().isoformat(),
            'status': 'pending',
            'case_data': case_data
        }
        
        # Update case with payment link
        case_data['payment_link'] = {
            'token': payment_token,
            'price': price,
            'created_at': datetime.now().isoformat(),
            'status': 'pending'
        }
        
        with open(case_file, 'w') as f:
            json.dump(case_data, f, indent=2)
        
        # Generate payment URL
        payment_url = f"/payment.html?token={payment_token}"
        
        return jsonify({
            'success': True,
            'payment_url': payment_url,
            'payment_token': payment_token,
            'full_url': f"https://turboresponsehq.ai{payment_url}"
        }), 200
        
    except Exception as e:
        print(f"Error generating payment link: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@automation_bp.route('/api/payment-link/<token>', methods=['GET'])
def get_payment_link_data(token):
    """Get case data for payment link"""
    try:
        if token not in payment_links:
            return jsonify({'success': False, 'error': 'Invalid or expired link'}), 404
        
        payment_data = payment_links[token]
        
        return jsonify({
            'success': True,
            'case': payment_data['case_data'],
            'price': payment_data['price'],
            'status': payment_data['status']
        }), 200
        
    except Exception as e:
        print(f"Error retrieving payment link: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@automation_bp.route('/api/payment-confirm/<token>', methods=['POST'])
def confirm_payment(token):
    """Confirm payment completion"""
    try:
        if token not in payment_links:
            return jsonify({'success': False, 'error': 'Invalid or expired link'}), 404
        
        payment_data = payment_links[token]
        case_id = payment_data['case_id']
        
        # Update payment link status
        payment_links[token]['status'] = 'paid'
        payment_links[token]['paid_at'] = datetime.now().isoformat()
        
        # Update case file
        case_file = os.path.join(AUTOMATION_CONFIG['data_storage_path'], f"{case_id}_submission.json")
        with open(case_file, 'r') as f:
            case_data = json.load(f)
        
        case_data['payment_status'] = 'approved'
        case_data['payment_approved_at'] = datetime.now().isoformat()
        case_data['payment_link']['status'] = 'paid'
        case_data['payment_link']['paid_at'] = datetime.now().isoformat()
        
        with open(case_file, 'w') as f:
            json.dump(case_data, f, indent=2)
        
        return jsonify({
            'success': True,
            'message': 'Payment confirmed',
            'case_id': case_id
        }), 200
        
    except Exception as e:
        print(f"Error confirming payment: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500



@automation_bp.route('/api/admin/mark-paid', methods=['POST'])
def mark_case_paid():
    """Manually mark a case as paid"""
    try:
        data = request.get_json()
        case_id = data.get('case_id')
        payment_method = data.get('payment_method', 'manual')
        
        if not case_id:
            return jsonify({
                'success': False,
                'error': 'case_id is required'
            }), 400
        
        data_path = AUTOMATION_CONFIG['data_storage_path']
        filepath = os.path.join(data_path, f"{case_id}_submission.json")
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'error': 'Case not found'
            }), 404
        
        # Read case data
        with open(filepath, 'r') as f:
            case_data = json.load(f)
        
        # Update payment status
        case_data['payment_status'] = 'paid'
        case_data['payment_method'] = payment_method
        case_data['payment_marked_at'] = datetime.now().isoformat()
        
        # Save updated data
        with open(filepath, 'w') as f:
            json.dump(case_data, f, indent=2)
        
        return jsonify({
            'success': True,
            'message': 'Case marked as paid',
            'case_id': case_id
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@automation_bp.route('/api/admin/case/<case_id>/archive', methods=['POST'])
def archive_case(case_id):
    """Archive a case (hide from main view)"""
    try:
        data_path = AUTOMATION_CONFIG['data_storage_path']
        filepath = os.path.join(data_path, f"{case_id}_submission.json")
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'error': 'Case not found'
            }), 404
        
        # Read case data
        with open(filepath, 'r') as f:
            case_data = json.load(f)
        
        # Update status to archived
        case_data['status'] = 'archived'
        case_data['archived_at'] = datetime.now().isoformat()
        
        # Save updated data
        with open(filepath, 'w') as f:
            json.dump(case_data, f, indent=2)
        
        return jsonify({
            'success': True,
            'message': 'Case archived successfully',
            'case_id': case_id
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@automation_bp.route('/api/admin/case/<case_id>/unarchive', methods=['POST'])
def unarchive_case(case_id):
    """Unarchive a case (restore to main view)"""
    try:
        data_path = AUTOMATION_CONFIG['data_storage_path']
        filepath = os.path.join(data_path, f"{case_id}_submission.json")
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'error': 'Case not found'
            }), 404
        
        # Read case data
        with open(filepath, 'r') as f:
            case_data = json.load(f)
        
        # Update status back to submitted
        case_data['status'] = 'submitted'
        case_data['unarchived_at'] = datetime.now().isoformat()
        
        # Save updated data
        with open(filepath, 'w') as f:
            json.dump(case_data, f, indent=2)
        
        return jsonify({
            'success': True,
            'message': 'Case unarchived successfully',
            'case_id': case_id
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@automation_bp.route('/api/admin/case/<case_id>/delete', methods=['DELETE'])
def delete_case(case_id):
    """Permanently delete a case and all associated documents"""
    try:
        data_path = AUTOMATION_CONFIG['data_storage_path']
        doc_path = AUTOMATION_CONFIG['document_storage_path']
        
        filepath = os.path.join(data_path, f"{case_id}_submission.json")
        case_doc_path = os.path.join(doc_path, case_id)
        
        # Check if case exists
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'error': 'Case not found'
            }), 404
        
        # Delete case file
        os.remove(filepath)
        
        # Delete associated documents if they exist
        if os.path.exists(case_doc_path):
            import shutil
            shutil.rmtree(case_doc_path)
        
        return jsonify({
            'success': True,
            'message': 'Case deleted permanently',
            'case_id': case_id
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500



@automation_bp.route('/api/paypal-payment-complete', methods=['POST'])
def paypal_payment_complete():
    """Handle PayPal payment completion"""
    try:
        data = request.get_json()
        payment_token = data.get('payment_token')
        order_id = data.get('order_id')
        payer_email = data.get('payer_email')
        payer_name = data.get('payer_name')
        amount = data.get('amount')
        
        if not payment_token or not order_id:
            return jsonify({
                'success': False,
                'error': 'Missing required fields'
            }), 400
        
        # Get case ID from payment_links
        if payment_token not in payment_links:
            return jsonify({
                'success': False,
                'error': 'Invalid payment token'
            }), 404
        
        case_id = payment_links[payment_token]['case_id']
        
        # Update payment link status
        payment_links[payment_token]['status'] = 'paid'
        payment_links[payment_token]['paid_at'] = datetime.now().isoformat()
        payment_links[payment_token]['paypal_order_id'] = order_id
        payment_links[payment_token]['payer_email'] = payer_email
        payment_links[payment_token]['payer_name'] = payer_name
        
        # Update case file
        data_path = AUTOMATION_CONFIG['data_storage_path']
        filepath = os.path.join(data_path, f"{case_id}_submission.json")
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'error': 'Case file not found'
            }), 404
        
        with open(filepath, 'r') as f:
            case_data = json.load(f)
        
        # Update case with payment info
        case_data['payment_status'] = 'paid'
        case_data['payment_method'] = 'paypal'
        case_data['payment_approved_at'] = datetime.now().isoformat()
        case_data['paypal_order_id'] = order_id
        case_data['paypal_payer_email'] = payer_email
        case_data['paypal_payer_name'] = payer_name
        case_data['payment_amount'] = amount
        
        if 'payment_link' in case_data:
            case_data['payment_link']['status'] = 'paid'
            case_data['payment_link']['paid_at'] = datetime.now().isoformat()
        
        with open(filepath, 'w') as f:
            json.dump(case_data, f, indent=2)
        
        return jsonify({
            'success': True,
            'message': 'Payment confirmed',
            'case_id': case_id,
            'order_id': order_id
        }), 200
        
    except Exception as e:
        print(f"Error processing PayPal payment: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500



def scan_case_documents(case_id):
    """
    Scan all uploaded documents for a case and extract text content.
    Supports PDF, images (with OCR), and Word documents.
    """
    try:
        import PyPDF2
        from docx import Document
        from PIL import Image
        import pytesseract
        from pdf2image import convert_from_path
        
        doc_path = AUTOMATION_CONFIG['document_storage_path']
        case_folder = os.path.join(doc_path, case_id)
        
        if not os.path.exists(case_folder):
            return "No documents uploaded."
        
        extracted_texts = []
        files = os.listdir(case_folder)
        
        if not files:
            return "No documents uploaded."
        
        for filename in files:
            filepath = os.path.join(case_folder, filename)
            file_ext = filename.lower().split('.')[-1]
            
            try:
                # PDF files
                if file_ext == 'pdf':
                    with open(filepath, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                        
                        if text.strip():
                            extracted_texts.append(f"\n--- Document: {filename} ---\n{text}")
                        else:
                            # If no text extracted, try OCR on PDF images
                            try:
                                images = convert_from_path(filepath)
                                ocr_text = ""
                                for i, image in enumerate(images):
                                    ocr_text += pytesseract.image_to_string(image) + "\n"
                                if ocr_text.strip():
                                    extracted_texts.append(f"\n--- Document: {filename} (OCR) ---\n{ocr_text}")
                            except Exception as ocr_error:
                                print(f"OCR failed for {filename}: {ocr_error}")
                
                # Image files (OCR)
                elif file_ext in ['jpg', 'jpeg', 'png', 'bmp', 'tiff']:
                    image = Image.open(filepath)
                    text = pytesseract.image_to_string(image)
                    if text.strip():
                        extracted_texts.append(f"\n--- Document: {filename} (Image OCR) ---\n{text}")
                
                # Word documents
                elif file_ext in ['doc', 'docx']:
                    doc = Document(filepath)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    if text.strip():
                        extracted_texts.append(f"\n--- Document: {filename} ---\n{text}")
                        
            except Exception as e:
                print(f"Error processing {filename}: {e}")
                continue
        
        if extracted_texts:
            return "\n".join(extracted_texts)
        else:
            return "Documents were uploaded but text extraction failed. Files may be images or scanned documents that require manual review."
            
    except Exception as e:
        print(f"Error scanning documents: {e}")
        return "Error scanning documents. Manual review required."


@automation_bp.route('/api/admin/case/<case_id>/ai-analysis', methods=['POST'])
def generate_ai_analysis(case_id):
    """Generate AI-powered case analysis using OpenAI"""
    try:
        import os
        from openai import OpenAI
        
        # Get OpenAI API key from environment
        api_key = os.environ.get('OPENAI_API_KEY')
        if not api_key:
            return jsonify({
                'success': False,
                'error': 'OpenAI API key not configured'
            }), 500
        
        # Load case data
        data_path = AUTOMATION_CONFIG['data_storage_path']
        filepath = os.path.join(data_path, f"{case_id}_submission.json")
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'error': 'Case not found'
            }), 404
        
        with open(filepath, 'r') as f:
            case_data = json.load(f)
        
        # Extract case information
        client_data = case_data.get('client_data', {})
        category = client_data.get('category', 'Unknown')
        description = client_data.get('caseDescription', 'No description provided')
        client_name = client_data.get('fullName', 'Unknown')
        
        # Scan uploaded documents and extract text
        document_text = scan_case_documents(case_id)
        
        # Load philosophy from organized sections
        philosophy_text = load_philosophy_sections()
        
        # Create AI prompt
        prompt = f"""You are the AI assistant for Turbo Response, a Consumer Rights Agency that empowers everyday citizens to fight back against powerful institutions.

{philosophy_text}

**CASE INFORMATION:**
- Client: {client_name}
- Category: {category}
- Description: {description}

**UPLOADED DOCUMENTS ANALYSIS:**
{document_text}

**YOUR ANALYSIS MUST:**

1. **CLASSIFY URGENCY TIER** (choose ONE):
   - **TIER 1 - EMERGENCY**: Panic/Fear state. Immediate threats (eviction, garnishment, IRS levy, bank freeze). Role: RESCUER. Tone: Calm + directive.
   - **TIER 2 - RECOVERY**: Frustration/Overwhelm state. Active disputes (debt collection, credit errors, fraud). Role: ADVOCATE. Tone: Confident + motivational.
   - **TIER 3 - REBUILDING**: Caution/Uncertainty state. Long-term repair (credit rebuilding, ChexSystems cleanup). Role: MENTOR. Tone: Encouraging + educational.
   - **TIER 4 - EMPOWERMENT**: Confidence/Growth state. Strategic growth (business setup, AI tools, education). Role: LEADER. Tone: Inspirational + strategic.

2. **DETECT LEGAL VIOLATIONS** (Primary focus):
   - FDCPA (Fair Debt Collection Practices Act) - §807, §809, §805
   - FCRA (Fair Credit Reporting Act) - §609, §611, §615
   - ECOA (Equal Credit Opportunity Act)
   - FCBA (Fair Credit Billing Act)
   - TCPA (Telephone Consumer Protection Act)
   - Fair Housing Act
   - IRS Taxpayer Bill of Rights
   → Identify specific violations and offensive strategies (file complaints with CFPB, FTC, AG)

3. **GAME PLAN** (5-8 action steps):
   - Start with immediate actions
   - Include offensive strategies (complaints, violations)
   - Show cost of inaction to create urgency
   - Be specific and actionable

4. **EMOTIONAL COMMUNICATION**:
   - Tier 1: "Take a breath — here's what we do now."
   - Tier 2: "You've got rights. Let's make them listen."
   - Tier 3: "Now we rebuild step-by-step."
   - Tier 4: "You control the narrative from here."
   - Always say: "You have rights — let's use them."
   - Never say: "You'll definitely win" or give legal advice

5. **PRICING LOGIC** (comprehensive calculation):
   
   **Step 1 - Base Price (Amount at Stake):**
   - Under $1,000: $149-$299
   - $1,000-$5,000: $299-$499
   - $5,000-$15,000: $499-$699
   - $15,000-$50,000: $699-$999
   - $50,000+: $999-$1,500+
   
   **Step 2 - Add Multipliers:**
   - Time/Retainer (ongoing case): +$200-$400
   - Court/Hearing scheduled: +$200-$300
   - Federal agencies (IRS, EEOC, HUD): +$150-$300
   - Multiple agencies: +$150-$300
   - Extensive documentation (10+ docs): +$100-$200
   - Urgent deadline (<30 days): +$100-$200
   - Legal violations (6+): +$200-$400
   
   **Step 3 - Calculate Total:**
   Start with base price from amount at stake, then add applicable multipliers.
   
   **Example:** IRS case, $23K owed, hearing scheduled:
   - Base: $699 (amount: $15K-$50K)
   - + Hearing: +$250
   - + IRS (federal): +$150
   - + Multiple deadlines: +$150
   - + Retainer likely: +$300
   - **Total: $1,549**
   
   **IMPORTANT:** Extract dollar amounts from documents and case description to determine base price.

6. **SERVICE RECOMMENDATION**:
   - Starter Plan: Single issue, quick resolution
   - Standard Plan: 2-3 issues, moderate complexity
   - Comprehensive Plan: Multiple issues, ongoing support

7. **TIMELINE ESTIMATE**:
   - Simple: 3-5 days
   - Moderate: 7-14 days
   - Complex: 15-30 days
   - Very Complex: 30+ days

8. **SUCCESS PROBABILITY**:
   - Base: 70% (adjust ± by evidence clarity, law strength, agency type)
   - Show what increases/decreases success

9. **KEY CHALLENGES**:
   - Identify obstacles
   - Show cost of waiting
   - Pre-empt objections

10. **REQUIRED EVIDENCE**:
    - What documents client must provide
    - What violations to document

Format your response as JSON with these exact keys:
{{
  \"tier\": \"TIER 1|TIER 2|TIER 3|TIER 4\",
  \"tier_label\": \"Emergency|Recovery|Rebuilding|Empowerment\",
  \"emotional_state\": \"Panic|Frustration|Caution|Confidence\",
  \"ai_role\": \"Rescuer|Advocate|Mentor|Leader\",
  \"complexity\": \"Simple|Moderate|Complex|Very Complex\",
  \"legal_violations\": [\"FDCPA §807\", \"FCRA §609\", ...],
  \"game_plan\": [\"step 1\", \"step 2\", ...],
  \"emotional_message\": \"Tier-appropriate message\",
  \"challenges\": [\"challenge 1\", \"challenge 2\", ...],
  \"required_evidence\": [\"doc 1\", \"doc 2\", ...],
  \"service_recommendation\": \"Starter|Standard|Comprehensive\",
  \"pricing_min\": 149,
  \"pricing_max\": 199,
  \"pricing_reason\": \"explanation\",
  \"timeline_days\": 7,
  \"success_probability\": 75,
  \"success_explanation\": \"brief explanation\",
  \"relevant_laws\": [\"FDCPA\", \"FCRA\", ...],
  \"urgency_message\": \"Cost of waiting explanation\"
}}"""
        
        # Call OpenAI API
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert legal assistant specializing in debt disputes and consumer rights. Always respond with valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            response_format={"type": "json_object"}
        )
        
        # Parse AI response
        ai_response = response.choices[0].message.content
        analysis = json.loads(ai_response)
        
        # Store analysis with case data
        case_data['ai_analysis'] = analysis
        case_data['ai_analysis_generated_at'] = datetime.now().isoformat()
        
        with open(filepath, 'w') as f:
            json.dump(case_data, f, indent=2)
        
        return jsonify({
            'success': True,
            'analysis': analysis,
            'case_id': case_id
        }), 200
        
    except Exception as e:
        print(f"Error generating AI analysis: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500




# Philosophy Management Endpoints
@automation_bp.route('/api/admin/get-philosophy', methods=['GET'])
def get_philosophy():
    """Get current AI philosophy"""
    try:
        philosophy_file = os.path.join(os.path.dirname(__file__), '..', 'data', 'philosophy.txt')
        
        if os.path.exists(philosophy_file):
            with open(philosophy_file, 'r') as f:
                philosophy = f.read()
        else:
            # Return default philosophy
            philosophy = """TURBO RESPONSE AI PHILOSOPHY

**MISSION:** Defend consumers from predatory or unlawful actions with speed + precision. "We don't just respond fast — we respond smart."

**CORE PHILOSOPHY:**
- Protect the consumer → Simplify the law → Drive resolution fast
- Lead with emotion → Support with logic → Close with empowerment
- Create urgency by showing the cost of waiting
- Fight back offensively, not just defensively
- "You're not buying a service — you're buying back your control"

**4-TIER SYSTEM:**
- TIER 1 - EMERGENCY: Panic/Fear state. Immediate threats. Role: RESCUER.
- TIER 2 - RECOVERY: Frustration/Overwhelm state. Active disputes. Role: ADVOCATE.
- TIER 3 - REBUILDING: Caution/Uncertainty state. Long-term repair. Role: MENTOR.
- TIER 4 - EMPOWERMENT: Confidence/Growth state. Strategic growth. Role: LEADER.

**PRICING LOGIC:**
- Tier 1 Emergency: $149-$199
- Tier 2 Recovery: $349-$499
- Tier 3 Rebuilding: $699-$999
- Tier 4 Empowerment: $999+

**LEGAL FOCUS:**
FDCPA, FCRA, ECOA, FCBA, TCPA, Fair Housing Act, IRS Taxpayer Bill of Rights

**CORE PRINCIPLES:**
1. Always identify legal violations first
2. Lead with emotion → follow with law
3. Show cost of inaction to create urgency
4. Protect client energy — remove confusion
5. Never promise results; promise effort + speed
6. Offer clarity over complexity
7. Teach rights with every interaction
"""
        
        return jsonify({'success': True, 'philosophy': philosophy}), 200
        
    except Exception as e:
        print(f"Error getting philosophy: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@automation_bp.route('/api/admin/save-philosophy', methods=['POST'])
def save_philosophy():
    """Save updated AI philosophy"""
    try:
        data = request.get_json()
        philosophy = data.get('philosophy', '')
        
        # Create data directory if it doesn't exist
        data_dir = os.path.join(os.path.dirname(__file__), '..', 'data')
        os.makedirs(data_dir, exist_ok=True)
        
        philosophy_file = os.path.join(data_dir, 'philosophy.txt')
        
        with open(philosophy_file, 'w') as f:
            f.write(philosophy)
        
        return jsonify({'success': True}), 200
        
    except Exception as e:
        print(f"Error saving philosophy: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@automation_bp.route('/api/admin/reset-philosophy', methods=['POST'])
def reset_philosophy():
    """Reset philosophy to default"""
    try:
        philosophy_file = os.path.join(os.path.dirname(__file__), '..', 'data', 'philosophy.txt')
        
        if os.path.exists(philosophy_file):
            os.remove(philosophy_file)
        
        return jsonify({'success': True}), 200
        
    except Exception as e:
        print(f"Error resetting philosophy: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500




# Section-based Philosophy Management
@automation_bp.route('/api/admin/get-philosophy-sections', methods=['GET'])
def get_philosophy_sections():
    """Get philosophy organized by sections"""
    try:
        data_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'philosophy_sections')
        sections = {}
        
        section_names = ['mission', 'tiers', 'pricing', 'legal', 'communication', 'website']
        
        for section in section_names:
            section_file = os.path.join(data_dir, f'{section}.txt')
            if os.path.exists(section_file):
                with open(section_file, 'r') as f:
                    sections[section] = f.read()
            else:
                # Load default content for each section
                sections[section] = get_default_section(section)
        
        return jsonify({'success': True, 'sections': sections}), 200
        
    except Exception as e:
        print(f"Error getting philosophy sections: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@automation_bp.route('/api/admin/save-philosophy-sections', methods=['POST'])
def save_philosophy_sections():
    """Save philosophy sections"""
    try:
        data = request.get_json()
        sections_data = data.get('sections', {})
        
        # Create sections directory
        data_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'philosophy_sections')
        os.makedirs(data_dir, exist_ok=True)
        
        # Save each section
        for section_name, content in sections_data.items():
            section_file = os.path.join(data_dir, f'{section_name}.txt')
            with open(section_file, 'w') as f:
                f.write(content)
        
        return jsonify({'success': True}), 200
        
    except Exception as e:
        print(f"Error saving philosophy sections: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


def get_default_section(section_name):
    """Get default content for a philosophy section"""
    defaults = {
        'mission': """**MISSION:** Defend consumers from predatory or unlawful actions with speed + precision. "We don't just respond fast — we respond smart."

**CORE VALUES:**
- Protect the consumer → Simplify the law → Drive resolution fast
- Lead with emotion → Support with logic → Close with empowerment
- Create urgency by showing the cost of waiting
- Fight back offensively, not just defensively
- "You're not buying a service — you're buying back your control" """,
        
        'tiers': """**4-TIER CLASSIFICATION SYSTEM:**

**TIER 1 - EMERGENCY:** Panic/Fear state. Immediate threats (eviction, wage garnishment, bank freeze). Role: RESCUER.
- Emotional State: Panic, fear, desperation
- Timeline: Immediate (24-48 hours)
- Pricing: $149-$199

**TIER 2 - RECOVERY:** Frustration/Overwhelm state. Active disputes (debt validation, credit disputes). Role: ADVOCATE.
- Emotional State: Frustration, anger, overwhelm
- Timeline: 7-14 days
- Pricing: $349-$499

**TIER 3 - REBUILDING:** Caution/Uncertainty state. Long-term repair (credit rebuilding, record cleanup). Role: MENTOR.
- Emotional State: Caution, uncertainty, hope
- Timeline: 30-90 days
- Pricing: $699-$999

**TIER 4 - EMPOWERMENT:** Confidence/Growth state. Strategic growth (business setup, automation). Role: LEADER.
- Emotional State: Confidence, curiosity, ambition
- Timeline: Ongoing
- Pricing: $999+""",
        
        'pricing': """**PRICING LOGIC:**

**Factors that increase price:**
- Multiple violations or creditors
- Court involvement or legal deadlines
- Complex documentation requirements
- High-stakes outcomes (home, job, large amounts)

**Factors that decrease price:**
- Simple single-issue cases
- Clear documentation provided
- Standard timeline (no rush)
- Straightforward resolution path

**Service Tiers:**
- Starter Plan: Single issue, standard timeline
- Standard Plan: Multiple issues, moderate complexity
- Comprehensive Plan: Full-service, urgent, complex""",
        
        'legal': """**LEGAL FOCUS AREAS:**

**FDCPA (Fair Debt Collection Practices Act):**
- §807: False, deceptive, or misleading representations
- §805: Communication restrictions
- §809: Validation of debts

**FCRA (Fair Credit Reporting Act):**
- §609: Disclosure requirements
- §611: Dispute procedures
- §616: Civil liability

**ECOA (Equal Credit Opportunity Act):**
- Discrimination in credit decisions

**TCPA (Telephone Consumer Protection Act):**
- Unwanted calls, texts, robocalls

**Fair Housing Act:**
- Housing discrimination

**IRS Taxpayer Bill of Rights:**
- Right to challenge IRS position
- Right to appeal decisions""",
        
        'communication': """**COMMUNICATION FRAMEWORK:**

**ERBN (Emotional Reasons to Buy Now):**
- Relief from stress and fear
- Protection and safety
- Control over situation
- Justice and validation
- Hope for resolution

**LRBN (Logical Reasons to Buy Now):**
- Time-sensitive legal windows
- Prevent escalation and additional costs
- Expert knowledge and efficiency
- Cost-effective vs. attorney fees
- Proven track record

**DRAB (Diminish, Reframe, Associate Pain, Bridge):**
- Diminish objections with empathy
- Reframe perspective to show value
- Associate pain with inaction
- Bridge to confident decision

**Tone Switching:**
- TIER 1: Urgent, protective, commanding (RESCUER)
- TIER 2: Supportive, strategic, empowering (ADVOCATE)
- TIER 3: Patient, educational, encouraging (MENTOR)
- TIER 4: Collaborative, visionary, inspiring (LEADER)""",
        
        'website': """**WEBSITE-SPECIFIC INSTRUCTIONS:**

**Key Features to Highlight:**
- AI-powered case analysis (instant evaluation)
- 24-hour response time
- PayPal payment option (instant confirmation, Pay in 4)
- Secure intake form at /intake
- Admin dashboard for case tracking

**Conversion Points:**
- Direct new clients to intake form
- Emphasize speed and technology advantage
- Highlight success stories and testimonials
- Show clear pricing and service tiers
- Offer free case evaluation

**Unique Selling Points:**
- "We don't just respond fast — we respond smart"
- AI-driven consumer defense
- Same power large organizations use
- Level the playing field"""
    }
    
    return defaults.get(section_name, '')




def load_philosophy_sections():
    """Load and combine all philosophy sections"""
    try:
        data_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'philosophy_sections')
        sections = []
        
        section_names = ['mission', 'tiers', 'pricing', 'legal', 'communication', 'website']
        
        for section in section_names:
            section_file = os.path.join(data_dir, f'{section}.txt')
            if os.path.exists(section_file):
                with open(section_file, 'r') as f:
                    content = f.read().strip()
                    if content:
                        sections.append(content)
            else:
                # Use default content
                default_content = get_default_section(section)
                if default_content:
                    sections.append(default_content)
        
        return '\n\n'.join(sections)
        
    except Exception as e:
        print(f"Error loading philosophy sections: {e}")
        # Return fallback philosophy
        return """**MISSION:** Defend consumers from predatory or unlawful actions with speed + precision.

**CORE PHILOSOPHY:**
- Protect the consumer → Simplify the law → Drive resolution fast
- Lead with emotion → Support with logic → Close with empowerment
- Create urgency by showing the cost of waiting"""




@automation_bp.route('/api/admin/upload-philosophy-document', methods=['POST'])
def upload_philosophy_document():
    """Upload and process document files (.txt, .docx, .pdf) for philosophy sections"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400
        
        file = request.files['file']
        section = request.form.get('section')
        
        if not section:
            return jsonify({'success': False, 'error': 'No section specified'}), 400
        
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        # Get file extension
        filename = file.filename.lower()
        
        # Extract text based on file type
        text_content = ''
        
        if filename.endswith('.txt'):
            # Plain text file
            text_content = file.read().decode('utf-8')
            
        elif filename.endswith('.docx'):
            # Word document
            try:
                from docx import Document
                import io
                doc = Document(io.BytesIO(file.read()))
                text_content = '\n\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            except Exception as e:
                return jsonify({'success': False, 'error': f'Error reading Word document: {str(e)}'}), 500
                
        elif filename.endswith('.pdf'):
            # PDF file
            try:
                from PyPDF2 import PdfReader
                import io
                pdf = PdfReader(io.BytesIO(file.read()))
                text_content = '\n\n'.join([page.extract_text() for page in pdf.pages])
            except Exception as e:
                return jsonify({'success': False, 'error': f'Error reading PDF: {str(e)}'}), 500
        else:
            return jsonify({'success': False, 'error': 'Unsupported file type. Please upload .txt, .docx, or .pdf files.'}), 400
        
        if not text_content.strip():
            return jsonify({'success': False, 'error': 'No text content found in file'}), 400
        
        # Save to section file
        data_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'philosophy_sections')
        os.makedirs(data_dir, exist_ok=True)
        
        section_file = os.path.join(data_dir, f'{section}.txt')
        
        # Append to existing content or create new
        if os.path.exists(section_file):
            with open(section_file, 'r') as f:
                existing_content = f.read()
            text_content = existing_content + '\n\n' + text_content
        
        with open(section_file, 'w') as f:
            f.write(text_content)
        
        return jsonify({
            'success': True,
            'content': text_content,
            'message': f'Document uploaded and processed successfully'
        }), 200
        
    except Exception as e:
        print(f"Error uploading document: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

